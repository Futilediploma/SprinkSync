"""
Unallocated Manpower Report PDF/DOCX/Excel Export
Generates a report of projects with required manpower not yet allocated.
"""

from fastapi import APIRouter, Depends, Response
from sqlalchemy.orm import Session
from datetime import datetime
import io
from database import get_db
from api.auth import get_current_active_user
import models

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import math
import os

from api.pdf_colors import COLORS

router = APIRouter(prefix="/api/export", tags=["export"])


class ManpowerNeedsPDF:
    """Professional Unallocated Manpower Report PDF Generator"""

    def __init__(self, page_size=letter):
        self.page_size = page_size
        self.width, self.height = page_size
        self.margin_left = 0.5 * inch
        self.margin_right = 0.5 * inch
        self.margin_top = 0.6 * inch
        self.margin_bottom = 0.8 * inch

        self.col_widths = {
            'project_name': 2.2 * inch,
            'project_number': 0.85 * inch,
            'customer': 1.5 * inch,
            'start': 0.8 * inch,
            'end': 0.8 * inch,
            'men_required': 0.75 * inch,
            'status': 0.75 * inch,
            'out_of_town': 0.85 * inch,
        }

        self.row_height = 18
        self.buffer = io.BytesIO()
        self.canvas = canvas.Canvas(self.buffer, pagesize=self.page_size)
        self.page_number = 0
        self.total_pages = 1

    def draw_header(self, run_date: str, logo_path: str = None):
        c = self.canvas
        y = self.height - self.margin_top + 20

        logo_width = 0
        if logo_path and os.path.exists(logo_path):
            try:
                logo_height = 30
                img = ImageReader(logo_path)
                img_w, img_h = img.getSize()
                logo_width = logo_height * (img_w / img_h)
                c.drawImage(logo_path, self.margin_left, y - 15,
                            width=logo_width, height=logo_height, preserveAspectRatio=True)
                logo_width += 10
            except Exception:
                logo_width = 0

        c.setFont("Helvetica-Bold", 14)
        c.setFillColor(COLORS['text_primary'])
        c.drawString(self.margin_left + logo_width, y, "Unallocated Manpower Report")

        c.setFont("Helvetica", 9)
        c.setFillColor(COLORS['text_secondary'])
        c.drawRightString(self.width - self.margin_right, y, f"Page {self.page_number} of {self.total_pages}")
        c.drawRightString(self.width - self.margin_right, y - 12, f"Run Date: {run_date}")
        c.drawString(self.margin_left + logo_width, y - 15, self.subtitle)

        c.setStrokeColor(COLORS['grid_line'])
        c.setLineWidth(1)
        c.line(self.margin_left, y - 25, self.width - self.margin_right, y - 25)

    def draw_table_header(self, y_pos: float):
        c = self.canvas
        header_height = 25

        c.setFillColor(COLORS['header_bg'])
        c.rect(self.margin_left, y_pos - header_height,
               self.width - self.margin_left - self.margin_right, header_height, fill=1, stroke=0)

        c.setFillColor(COLORS['header_text'])
        c.setFont("Helvetica-Bold", 8)
        x = self.margin_left + 4
        text_y = y_pos - header_height / 2 - 3

        headers = [
            ('Project Name', self.col_widths['project_name']),
            ('Project #', self.col_widths['project_number']),
            ('Customer', self.col_widths['customer']),
            ('Start', self.col_widths['start']),
            ('End', self.col_widths['end']),
            ('Men Req.', self.col_widths['men_required']),
            ('Status', self.col_widths['status']),
            ('Out of Town', self.col_widths['out_of_town']),
        ]
        for header, width in headers:
            c.drawString(x, text_y, header)
            x += width

        return y_pos - header_height

    def draw_data_row(self, y_pos: float, project, row_index: int):
        c = self.canvas

        bg_color = COLORS['row_alt'] if row_index % 2 == 0 else COLORS['row_normal']
        c.setFillColor(bg_color)
        c.rect(self.margin_left, y_pos - self.row_height,
               self.width - self.margin_left - self.margin_right,
               self.row_height, fill=1, stroke=0)

        c.setStrokeColor(COLORS['grid_line'])
        c.setLineWidth(0.3)
        c.line(self.margin_left, y_pos - self.row_height,
               self.width - self.margin_right, y_pos - self.row_height)

        c.setFillColor(COLORS['text_primary'])
        c.setFont("Helvetica", 8)
        x = self.margin_left + 4
        text_y = y_pos - self.row_height + 5

        def truncate(s, col_key):
            s = str(s) if s else '—'
            max_chars = int(self.col_widths[col_key] / 4.5)
            return s[:max_chars - 2] + '..' if len(s) > max_chars else s

        c.drawString(x, text_y, truncate(project.name, 'project_name'))
        x += self.col_widths['project_name']

        c.drawString(x, text_y, truncate(project.project_number or '—', 'project_number'))
        x += self.col_widths['project_number']

        c.drawString(x, text_y, truncate(project.customer_name or '—', 'customer'))
        x += self.col_widths['customer']

        c.drawString(x, text_y, project.start_date.strftime('%d-%b-%y') if project.start_date else '—')
        x += self.col_widths['start']

        c.drawString(x, text_y, project.end_date.strftime('%d-%b-%y') if project.end_date else '—')
        x += self.col_widths['end']

        c.drawRightString(x + self.col_widths['men_required'] - 4, text_y,
                          str(project.required_manpower or 0))
        x += self.col_widths['men_required']

        c.drawString(x, text_y, project.status.capitalize())
        x += self.col_widths['status']

        c.drawString(x, text_y, 'Yes' if project.is_out_of_town else '')

        return y_pos - self.row_height

    def draw_summary_row(self, y_pos: float, total_men: int):
        c = self.canvas
        c.setFillColor(COLORS['subheader_bg'])
        c.rect(self.margin_left, y_pos - self.row_height,
               self.width - self.margin_left - self.margin_right,
               self.row_height, fill=1, stroke=0)

        c.setFillColor(COLORS['header_text'])
        c.setFont("Helvetica-Bold", 9)
        text_y = y_pos - self.row_height + 5

        c.drawString(self.margin_left + 4, text_y,
                     f"TOTAL — {total_men} men required across {self._project_count} projects")

        return y_pos - self.row_height

    def draw_footer(self):
        c = self.canvas
        y = self.margin_bottom - 20
        c.setStrokeColor(COLORS['grid_line'])
        c.setLineWidth(0.5)
        c.line(self.margin_left, y + 15, self.width - self.margin_right, y + 15)
        c.setFont("Helvetica", 8)
        c.setFillColor(COLORS['text_secondary'])
        c.drawCentredString(self.width / 2, y, "BFPE International — Unallocated Manpower Report")

    def generate(self, projects: list, subtitle: str = "Projects requiring manpower assignment"):
        self.subtitle = subtitle
        run_date = datetime.now().strftime('%B %d, %Y')
        logo_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
                                 'frontend', 'bfpe_logo.png')

        self._project_count = len(projects)
        total_men = sum(p.required_manpower or 0 for p in projects)

        rows_per_page = int((self.height - self.margin_top - self.margin_bottom - 80) / self.row_height)
        self.total_pages = max(1, math.ceil((len(projects) + 1) / rows_per_page))

        self.page_number = 1
        self.canvas.setTitle("Unallocated Manpower Report")
        self.draw_header(run_date, logo_path)

        content_top = self.height - self.margin_top - 30
        y = self.draw_table_header(content_top)

        for i, project in enumerate(projects):
            if y - self.row_height < self.margin_bottom + 40:
                self.draw_footer()
                self.canvas.showPage()
                self.page_number += 1
                self.draw_header(run_date, logo_path)
                y = self.draw_table_header(content_top)

            y = self.draw_data_row(y, project, i)

        y = self.draw_summary_row(y, total_men)
        self.draw_footer()
        self.canvas.save()
        self.buffer.seek(0)
        return self.buffer


@router.get("/pdf/manpower-needs")
def export_manpower_needs_pdf(
    project_ids: str = None,
    status_filter: str = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_active_user)
):
    """Export unallocated manpower report as a professional PDF."""
    query = db.query(models.Project).filter(
        models.Project.status.in_(['active', 'prospective']),
        models.Project.required_manpower > 0,
        models.Project.manpower_allocated == False
    )
    if project_ids:
        ids = [int(i) for i in project_ids.split(',') if i.strip().isdigit()]
        if ids:
            query = query.filter(models.Project.id.in_(ids))
    projects = query.order_by(models.Project.status, models.Project.name).all()

    if status_filter == 'active':
        subtitle = "Projects requiring manpower assignment — active jobs only"
    elif status_filter == 'prospective':
        subtitle = "Projects requiring manpower assignment — prospective jobs only"
    else:
        subtitle = "Projects requiring manpower assignment — active and prospective jobs"

    pdf_generator = ManpowerNeedsPDF(page_size=letter)
    pdf_buffer = pdf_generator.generate(projects=projects, subtitle=subtitle)

    return Response(
        pdf_buffer.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=Unallocated_Manpower_Report.pdf"}
    )


def _get_manpower_projects(db: Session, project_ids: str = None):
    query = db.query(models.Project).filter(
        models.Project.status.in_(['active', 'prospective']),
        models.Project.required_manpower > 0,
        models.Project.manpower_allocated == False
    )
    if project_ids:
        ids = [int(i) for i in project_ids.split(',') if i.strip().isdigit()]
        if ids:
            query = query.filter(models.Project.id.in_(ids))
    return query.order_by(models.Project.status, models.Project.name).all()


@router.get("/docx/manpower-needs")
def export_manpower_needs_docx(
    project_ids: str = None,
    status_filter: str = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_active_user)
):
    """Export unallocated manpower report as a Word document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    projects = _get_manpower_projects(db, project_ids)

    if status_filter == 'active':
        subtitle = "Projects requiring manpower assignment — active jobs only"
    elif status_filter == 'prospective':
        subtitle = "Projects requiring manpower assignment — prospective jobs only"
    else:
        subtitle = "Projects requiring manpower assignment — active and prospective jobs"

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("Unallocated Manpower Report")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(subtitle)
    sub_run.font.size = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    # Run date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Run Date: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()

    # Table
    headers = ['Project Name', 'Project #', 'Customer', 'Start', 'End', 'Men Req.', 'Status', 'Out of Town']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1e3a5f')
        tcPr.append(shd)

    # Data rows
    total_men = 0
    for i, project in enumerate(projects):
        row_cells = table.add_row().cells
        total_men += project.required_manpower or 0
        values = [
            project.name or '',
            project.project_number or '—',
            project.customer_name or '—',
            project.start_date.strftime('%d-%b-%y') if project.start_date else '—',
            project.end_date.strftime('%d-%b-%y') if project.end_date else '—',
            str(project.required_manpower or 0),
            project.status.capitalize(),
            'Yes' if project.is_out_of_town else '',
        ]
        fill = 'f0f4f8' if i % 2 == 0 else 'ffffff'
        for j, val in enumerate(values):
            row_cells[j].text = val
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(9)
            tc = row_cells[j]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)

    # Summary row
    summary_row = table.add_row().cells
    summary_row[0].merge(summary_row[-1])
    summary_row[0].text = f"TOTAL — {total_men} men required across {len(projects)} projects"
    summary_row[0].paragraphs[0].runs[0].bold = True
    summary_row[0].paragraphs[0].runs[0].font.size = Pt(10)
    summary_row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    tc = summary_row[0]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1e3a5f')
    tcPr.append(shd)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Unallocated_Manpower_Report.docx"}
    )


@router.get("/excel/manpower-needs")
def export_manpower_needs_excel(
    project_ids: str = None,
    status_filter: str = None,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_active_user)
):
    """Export unallocated manpower report as an Excel spreadsheet."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    projects = _get_manpower_projects(db, project_ids)

    if status_filter == 'active':
        subtitle = "Active jobs only"
    elif status_filter == 'prospective':
        subtitle = "Prospective jobs only"
    else:
        subtitle = "Active and prospective jobs"

    wb = Workbook()
    ws = wb.active
    ws.title = "Unallocated Manpower"

    # Title rows
    ws.merge_cells('A1:H1')
    ws['A1'] = "Unallocated Manpower Report"
    ws['A1'].font = Font(bold=True, size=14, color='1e3a5f')

    ws.merge_cells('A2:H2')
    ws['A2'] = f"Projects requiring manpower assignment — {subtitle}"
    ws['A2'].font = Font(size=9, color='6b7280')

    ws.merge_cells('A3:H3')
    ws['A3'] = f"Run Date: {datetime.now().strftime('%B %d, %Y')}"
    ws['A3'].font = Font(size=9, color='6b7280')

    # Header row
    headers = ['Project Name', 'Project #', 'Customer', 'Start Date', 'End Date', 'Men Required', 'Status', 'Out of Town']
    header_fill = PatternFill(start_color='1e3a5f', end_color='1e3a5f', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF', size=10)
    thin = Side(style='thin', color='d1d5db')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    ws.row_dimensions[5].height = 20

    # Data rows
    alt_fill = PatternFill(start_color='f0f4f8', end_color='f0f4f8', fill_type='solid')
    total_men = 0

    for i, project in enumerate(projects):
        row = 6 + i
        total_men += project.required_manpower or 0
        values = [
            project.name,
            project.project_number or '',
            project.customer_name or '',
            project.start_date.strftime('%d-%b-%y') if project.start_date else '',
            project.end_date.strftime('%d-%b-%y') if project.end_date else '',
            project.required_manpower or 0,
            project.status.capitalize(),
            'Yes' if project.is_out_of_town else '',
        ]
        fill = alt_fill if i % 2 == 0 else None
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = Font(size=9)
            cell.border = border
            cell.alignment = Alignment(vertical='center')
            if fill:
                cell.fill = fill

    # Summary row
    summary_row = 6 + len(projects)
    ws.merge_cells(f'A{summary_row}:H{summary_row}')
    ws[f'A{summary_row}'] = f"TOTAL — {total_men} men required across {len(projects)} projects"
    ws[f'A{summary_row}'].font = Font(bold=True, color='FFFFFF', size=10)
    ws[f'A{summary_row}'].fill = PatternFill(start_color='1e3a5f', end_color='1e3a5f', fill_type='solid')
    ws[f'A{summary_row}'].alignment = Alignment(vertical='center')
    ws.row_dimensions[summary_row].height = 20

    # Column widths
    col_widths = [35, 12, 22, 12, 12, 14, 12, 12]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    return Response(
        buffer.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=Unallocated_Manpower_Report.xlsx"}
    )
