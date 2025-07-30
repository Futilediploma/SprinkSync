"""
Document processing service for extracting project data from uploaded files
"""
import io
import json
import re
from typing import Dict, Any, List, Optional
from datetime import datetime
import PyPDF2
import docx
from openai import OpenAI
import os
from pathlib import Path

class DocumentProcessor:
    def __init__(self):
        # Initialize OpenAI client - you'll need to set OPENAI_API_KEY environment variable
        self.openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
    async def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Process uploaded document and extract project information
        """
        try:
            # Extract text based on file type
            extracted_text = await self._extract_text(file_content, filename)
            
            if not extracted_text.strip():
                return {"error": "No text could be extracted from the document"}
            
            # Use AI to extract structured project data
            project_data = await self._extract_project_data_with_ai(extracted_text, filename)
            
            # Add metadata
            project_data["extraction_metadata"] = {
                "source_filename": filename,
                "extracted_at": datetime.now().isoformat(),
                "text_length": len(extracted_text),
                "confidence_score": self._calculate_confidence(project_data)
            }
            
            return project_data
            
        except Exception as e:
            return {"error": f"Error processing document: {str(e)}"}
    
    async def _extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF or Word documents"""
        file_extension = Path(filename).suffix.lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf_text(file_content)
        elif file_extension in ['.docx', '.doc']:
            return self._extract_word_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def _extract_word_text(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting Word text: {str(e)}")
    
    async def _extract_project_data_with_ai(self, text: str, filename: str) -> Dict[str, Any]:
        """Use OpenAI to extract structured project data from text"""
        
        # Determine document type for better prompt engineering
        doc_type = self._identify_document_type(text, filename)
        
        prompt = self._build_extraction_prompt(text, doc_type)
        
        try:
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert construction project data extractor. Extract project information from documents and return valid JSON only."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            # Parse the JSON response
            ai_response = response.choices[0].message.content.strip()
            
            # Clean up the response to ensure it's valid JSON
            ai_response = self._clean_json_response(ai_response)
            
            extracted_data = json.loads(ai_response)
            
            # Post-process and validate the extracted data
            return self._post_process_extracted_data(extracted_data, doc_type)
            
        except Exception as e:
            # Fallback to rule-based extraction if AI fails
            return self._fallback_extraction(text)
    
    def _identify_document_type(self, text: str, filename: str) -> str:
        """Identify the type of construction document"""
        text_lower = text.lower()
        filename_lower = filename.lower()
        
        # Check for contract indicators
        contract_keywords = ['contract', 'agreement', 'scope of work', 'payment terms', 'change order']
        if any(keyword in text_lower for keyword in contract_keywords):
            return "contract"
        
        # Check for permit indicators
        permit_keywords = ['permit', 'building department', 'zoning', 'inspection']
        if any(keyword in text_lower for keyword in permit_keywords):
            return "permit"
        
        # Check for proposal indicators
        proposal_keywords = ['proposal', 'bid', 'estimate', 'quotation']
        if any(keyword in text_lower for keyword in proposal_keywords):
            return "proposal"
        
        # Check for insurance indicators
        insurance_keywords = ['insurance', 'liability', 'coverage', 'policy']
        if any(keyword in text_lower for keyword in insurance_keywords):
            return "insurance"
        
        return "general"
    
    def _build_extraction_prompt(self, text: str, doc_type: str) -> str:
        """Build AI prompt based on document type"""
        
        base_fields = """
        {
            "name": "Project name or title",
            "description": "Project description or scope",
            "job_number": "Job or project number",
            "job_type": "Type of construction work",
            "job_category": "Category or classification",
            "job_priority": "Priority level (High/Medium/Low)",
            "address": "Project address",
            "city": "City",
            "state": "State",
            "zip_code": "ZIP code",
            "county": "County",
            "client_name": "Client or owner name",
            "client_contact": "Client contact person",
            "client_phone": "Client phone number",
            "client_email": "Client email",
            "billing_address": "Billing address",
            "start_date": "Project start date (YYYY-MM-DD format)",
            "end_date": "Project end date (YYYY-MM-DD format)",
            "estimated_duration": "Duration in days (number only)",
            "contract_amount": "Contract amount (number only, no currency symbols)",
            "budget": "Project budget (number only)",
            "estimated_cost": "Estimated cost (number only)",
            "payment_terms": "Payment terms",
            "project_manager": "Project manager name",
            "site_supervisor": "Site supervisor name",
            "foreman": "Foreman name",
            "safety_officer": "Safety officer name",
            "permits_required": "Required permits",
            "permit_status": "Status of permits",
            "insurance_requirements": "Insurance requirements",
            "special_requirements": "Special requirements or notes",
            "safety_plan_required": "Is safety plan required (true/false)",
            "hazard_analysis": "Hazard analysis notes",
            "ppe_requirements": "PPE requirements"
        }
        """
        
        if doc_type == "contract":
            specific_instruction = "Focus on extracting contract terms, payment schedules, scope of work, and client information."
        elif doc_type == "permit":
            specific_instruction = "Focus on permit requirements, compliance items, inspection schedules, and regulatory information."
        elif doc_type == "proposal":
            specific_instruction = "Focus on project scope, cost estimates, timeline, and client requirements."
        else:
            specific_instruction = "Extract any relevant project information found in the document."
        
        prompt = f"""
        Extract construction project information from the following document text. {specific_instruction}
        
        Return ONLY a valid JSON object with the following structure. Use null for any fields that cannot be determined from the document:
        
        {base_fields}
        
        Additional guidelines:
        - For dates, use YYYY-MM-DD format or null if not found
        - For monetary values, extract only numbers (no $ signs or commas)
        - For boolean fields, use true/false or null
        - For text fields, extract exact text or reasonable interpretations
        - If information is unclear, use null rather than guessing
        
        Document text:
        {text[:4000]}  # Limit text to avoid token limits
        
        JSON response:
        """
        
        return prompt
    
    def _clean_json_response(self, response: str) -> str:
        """Clean up AI response to ensure valid JSON"""
        # Remove any markdown formatting
        response = re.sub(r'```json\s*', '', response)
        response = re.sub(r'```\s*$', '', response)
        
        # Find the JSON object
        start = response.find('{')
        end = response.rfind('}') + 1
        
        if start != -1 and end != 0:
            response = response[start:end]
        
        return response
    
    def _post_process_extracted_data(self, data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Post-process extracted data for validation and formatting"""
        processed = {}
        
        for key, value in data.items():
            if value is None or value == "" or value == "null":
                processed[key] = None
            elif key in ['contract_amount', 'budget', 'estimated_cost']:
                # Clean monetary values
                if isinstance(value, str):
                    cleaned = re.sub(r'[^\d.]', '', value)
                    processed[key] = float(cleaned) if cleaned else None
                else:
                    processed[key] = float(value) if value else None
            elif key == 'estimated_duration':
                # Clean duration to integer
                if isinstance(value, str):
                    cleaned = re.sub(r'[^\d]', '', value)
                    processed[key] = int(cleaned) if cleaned else None
                else:
                    processed[key] = int(value) if value else None
            elif key == 'safety_plan_required':
                # Convert to boolean
                if isinstance(value, str):
                    processed[key] = value.lower() in ['true', 'yes', '1', 'required']
                else:
                    processed[key] = bool(value) if value is not None else None
            elif key in ['start_date', 'end_date']:
                # Validate date format
                processed[key] = self._parse_date(value)
            else:
                processed[key] = value
        
        return processed
    
    def _parse_date(self, date_str: Any) -> Optional[str]:
        """Parse and validate date strings"""
        if not date_str or date_str == "null":
            return None
        
        try:
            # Try to parse various date formats
            if isinstance(date_str, str):
                # Try YYYY-MM-DD format first
                if re.match(r'\d{4}-\d{2}-\d{2}', date_str):
                    return date_str
                
                # Try other common formats and convert
                from dateutil import parser
                parsed_date = parser.parse(date_str)
                return parsed_date.strftime('%Y-%m-%d')
            
            return None
        except:
            return None
    
    def _fallback_extraction(self, text: str) -> Dict[str, Any]:
        """Fallback rule-based extraction if AI fails"""
        data = {}
        
        # Simple regex patterns for common fields
        patterns = {
            'client_name': r'(?:client|owner|customer):\s*([^\n]+)',
            'project_manager': r'(?:project manager|pm):\s*([^\n]+)',
            'address': r'(?:address|location):\s*([^\n]+)',
            'phone': r'(?:phone|tel):\s*([0-9\-\(\)\s]+)',
            'email': r'([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})'
        }
        
        for field, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                data[field] = match.group(1).strip()
        
        return data
    
    def _calculate_confidence(self, data: Dict[str, Any]) -> float:
        """Calculate confidence score based on extracted data completeness"""
        total_fields = 30  # Total number of possible fields
        filled_fields = sum(1 for value in data.values() if value is not None and value != "")
        
        return round(filled_fields / total_fields, 2)
